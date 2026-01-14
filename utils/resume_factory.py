"""
Helpers for generating small resume files for upload tests.

Why:
- The JobsNProfiles resume upload rejects files when the email parsed from the resume
  doesn't match the logged-in profile email.
- Shipping a static PDF makes tests brittle (email inside PDF may not match the test account).
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Union


PathLike = Union[str, Path]


_INVALID_FILENAME_CHARS = re.compile(r"[^a-zA-Z0-9._\-]+")


def _safe_filename(stem: str) -> str:
    stem = (stem or "").strip()
    if not stem:
        stem = "resume"
    stem = _INVALID_FILENAME_CHARS.sub("_", stem)
    stem = stem.strip("._-")
    return stem or "resume"


def create_rtf_resume(
    out_dir: PathLike,
    *,
    full_name: str,
    email: str,
    phone: str = "",
    filename: Optional[str] = None,
    extra_text: str = "",
) -> str:
    """
    Create a minimal .rtf resume file containing the provided contact details.
    Returns the absolute file path.
    """
    out_path = Path(out_dir).expanduser().resolve()
    out_path.mkdir(parents=True, exist_ok=True)

    email_norm = (email or "").strip()
    if not email_norm:
        raise ValueError("email is required to generate resume content")

    if not filename:
        filename = f"{_safe_filename(email_norm)}_auto_resume.rtf"
    if not filename.lower().endswith(".rtf"):
        filename = f"{filename}.rtf"

    file_path = (out_path / filename).resolve()

    # Keep email as plain text so server-side parsers can reliably extract it.
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    body_lines = [
        f"{full_name}".strip() or "Job Seeker",
        f"Email: {email_norm}",
        f"Phone: {phone}".strip() if phone else "",
        "",
        "Summary:",
        "Automation-generated resume for upload testing.",
        f"Generated at: {now}",
    ]
    if extra_text:
        body_lines.append("")
        body_lines.append(extra_text.strip())

    # Very small RTF; keep characters ASCII-safe.
    # Note: In RTF, backslashes and braces must be escaped.
    def esc(s: str) -> str:
        return (
            (s or "")
            .replace("\\", r"\\")
            .replace("{", r"\{")
            .replace("}", r"\}")
        )

    rtf = "{\\rtf1\\ansi\\deff0\n"
    for line in body_lines:
        if line is None:
            continue
        line = line.strip()
        if line == "":
            rtf += "\\par\n"
        else:
            rtf += f"{esc(line)}\\par\n"
    rtf += "}\n"

    file_path.write_text(rtf, encoding="utf-8")
    return str(file_path)


def create_docx_resume(
    out_dir: PathLike,
    *,
    full_name: str,
    email: str,
    phone: str = "",
    filename: Optional[str] = None,
    extra_text: str = "",
) -> str:
    """
    Create a minimal .docx resume containing the provided contact details.
    Returns the absolute file path.
    """
    # Import lazily so the rest of the suite doesn't hard-fail if dependency is missing.
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "python-docx is required to generate .docx resumes. Install it (pip install python-docx)."
        ) from e

    out_path = Path(out_dir).expanduser().resolve()
    out_path.mkdir(parents=True, exist_ok=True)

    email_norm = (email or "").strip()
    if not email_norm:
        raise ValueError("email is required to generate resume content")

    if not filename:
        filename = f"{_safe_filename(email_norm)}_auto_resume.docx"
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    file_path = (out_path / filename).resolve()

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc = Document()
    doc.add_heading(full_name.strip() or "Job Seeker", level=1)
    # Put the email as plain text on its own line (some parsers miss "Email: <addr>").
    doc.add_paragraph(email_norm)
    doc.add_paragraph(f"Email: {email_norm}")
    if phone:
        doc.add_paragraph(f"Phone: {phone.strip()}")
    doc.add_paragraph("")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph("Automation-generated resume for upload testing.")
    doc.add_paragraph(f"Generated at: {now}")
    if extra_text:
        doc.add_paragraph("")
        doc.add_paragraph(extra_text.strip())

    doc.save(str(file_path))
    return str(file_path)

